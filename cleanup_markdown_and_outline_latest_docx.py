from pathlib import Path
import re
import shutil

from docx import Document
from docx.oxml.ns import qn


WORK = Path(r"C:\chen_bauer_2024\docx_work")
SRC = WORK / "outline_cleanup_latest_input.docx"
LOCAL_OUT = WORK / "outline_cleanup_latest_final.docx"
FINAL_OUT = Path(r"N:\개인\다음논문\LLM_DIF_full_manuscript_draft_psychometric_style_Gemini범위수정_투고압축본_최종_탐색창정리.docx")

MAIN_HEADING = re.compile(r"^[1-4]\.\s+\S+")
SUB_HEADING = re.compile(r"^[1-4]\.\d+\s+\S+")


def remove_direct_outline(paragraph):
    ppr = paragraph._p.pPr
    if ppr is None:
        return
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is not None:
        ppr.remove(outline)


def clean_markdown_runs(paragraph):
    if not paragraph.runs:
        return
    full = "".join(run.text for run in paragraph.runs)
    new = full.replace("```", "").replace("**", "").replace("__", "").replace("`", "")
    # In this manuscript, remaining asterisks are markdown italics in references.
    new = new.replace("*", "")
    if new == full:
        return
    first = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first.text = new


def classify_heading(text):
    if text in {"국문초록", "English Abstract", "References"}:
        return "Heading 1"
    if text.startswith("부록 A.") or text.startswith("부록 B."):
        return "Heading 2"
    if bool(MAIN_HEADING.match(text)) and not bool(SUB_HEADING.match(text)):
        return "Heading 1"
    if bool(SUB_HEADING.match(text)):
        return "Heading 2"
    return None


doc = Document(SRC)
style_names = {s.name for s in doc.styles}

for paragraph in doc.paragraphs:
    clean_markdown_runs(paragraph)
    text = paragraph.text.strip()
    remove_direct_outline(paragraph)

    if text.startswith(".3.6 "):
        paragraph.text = paragraph.text.replace(".3.6 ", "3.6 ", 1)
        text = paragraph.text.strip()

    heading_style = classify_heading(text)
    if heading_style:
        paragraph.style = doc.styles[heading_style]
    elif text.startswith("표 "):
        paragraph.style = doc.styles["Caption"] if "Caption" in style_names else doc.styles["Normal"]
    else:
        paragraph.style = doc.styles["Normal"]

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                clean_markdown_runs(paragraph)
                remove_direct_outline(paragraph)

doc.save(LOCAL_OUT)
shutil.copy2(LOCAL_OUT, FINAL_OUT)
print(LOCAL_OUT)
print(FINAL_OUT)
