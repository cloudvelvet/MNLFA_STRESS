from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "proposal" / "RSDI_dissertation_proposal_ko_full_v2.md"
TARGET = ROOT / "docs" / "proposal" / "RSDI_dissertation_proposal_ko_full_v2.docx"

FONT = "Malgun Gothic"
MONO = "Consolas"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
LIGHT_FILL = "F4F6F9"


def set_run_font(run, *, name=FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
        ("Heading 4", 11, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def add_inline(paragraph, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, italic=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=MONO, size=9.5)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run)


def add_cover(doc):
    title = "종단 순서형 패널자료에서 잠재 변화와 측정기능 변화를 분해하는\n베이지안 종단 MNLFA 프레임워크"
    subtitle = "A Bayesian Longitudinal MNLFA Framework for Decomposing Latent Change and Measurement-Function Change in Ordinal Panel Data"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, size=24, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    set_run_font(run, size=13, italic=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(70)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("박사학위논문 연구계획서\nThree-Paper Dissertation Proposal")
    set_run_font(run, size=15, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("연구자: 이창현")
    set_run_font(run, size=12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026년 7월")
    set_run_font(run, size=11, color=MUTED)

    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows, i


def table_widths(cols):
    if cols == 2:
        return [2.0, 4.5]
    if cols == 3:
        return [1.5, 2.5, 2.5]
    if cols == 4:
        return [1.45, 1.75, 1.6, 1.7]
    if cols == 5:
        return [1.1, 1.35, 1.35, 1.35, 1.35]
    return [6.5 / cols] * cols


def add_table(doc, rows):
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_widths(table, table_widths(cols))
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            add_inline(p, text)
            for run in p.runs:
                set_run_font(run, size=9, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    configure_styles(doc)
    add_cover(doc)

    header = section.header.paragraphs[0]
    header.text = "박사학위논문 연구계획서 | RSDI"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED)
    set_page_number(section.footer.paragraphs[0])

    in_code = False
    code_buffer = []
    skipped_title = 0
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.right_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.05
                run = p.add_run("\n".join(code_buffer))
                set_run_font(run, name=MONO, size=8.5)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(raw)
            i += 1
            continue
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("# ") and skipped_title == 0:
            skipped_title += 1
            i += 1
            continue
        if line.startswith("## A Bayesian") and skipped_title == 1:
            skipped_title += 1
            i += 1
            continue
        if line.startswith("**박사학위논문 연구계획서"):
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading:
            level = min(len(heading.group(1)), 4)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2))
            i += 1
            continue
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.15)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            add_inline(p, line[2:])
            for run in p.runs:
                run.font.color.rgb = DARK_BLUE
            i += 1
            continue
        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", line))
            i += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", line)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, numbered.group(2))
            i += 1
            continue
        p = doc.add_paragraph()
        if line.startswith("**주제어:") or line.startswith("**Keywords:"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if line.startswith("Bauer,") or line.startswith("Chen,") or line.startswith("Kim,") or line.startswith("King-") or line.startswith("Liu,") or line.startswith("Oort,") or line.startswith("Padgett,") or line.startswith("Sprangers,") or line.startswith("Wallin,") or line.startswith("Widaman,") or line.startswith("이창현,"):
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
        add_inline(p, line)
        i += 1

    doc.core_properties.title = "RSDI 박사학위논문 연구계획서"
    doc.core_properties.subject = "Bayesian longitudinal ordinal MNLFA and RSDI"
    doc.core_properties.author = "이창현"
    doc.save(TARGET)
    print(TARGET)


if __name__ == "__main__":
    build()
