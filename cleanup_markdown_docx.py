from pathlib import Path
import shutil

from docx import Document


WORK = Path(r"C:\chen_bauer_2024\docx_work")
SRC = WORK / "markdown_cleanup_input.docx"
LOCAL_OUT = WORK / "markdown_cleanup_final.docx"
FINAL_OUT = Path(r"N:\개인\다음논문\LLM_DIF_full_manuscript_draft_psychometric_style_Gemini범위수정_투고압축본_최종_마크다운정리.docx")


MARKDOWN_TOKENS = [
    "```",
    "**",
    "__",
    "`",
]


def clean_text(text: str) -> str:
    new = text
    for token in MARKDOWN_TOKENS:
        new = new.replace(token, "")
    # Remaining single asterisks are only markdown italics in this manuscript's
    # references. Remove the marker while preserving the citation text.
    new = new.replace("*", "")
    return new


def clean_paragraph(paragraph):
    if not paragraph.runs:
        return
    full = "".join(run.text for run in paragraph.runs)
    new = clean_text(full)
    if new == full:
        return
    first = paragraph.runs[0]
    for run in paragraph.runs[1:]:
        run.text = ""
    first.text = new


doc = Document(SRC)

for paragraph in doc.paragraphs:
    clean_paragraph(paragraph)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                clean_paragraph(paragraph)

doc.save(LOCAL_OUT)
shutil.copy2(LOCAL_OUT, FINAL_OUT)
print(LOCAL_OUT)
print(FINAL_OUT)
