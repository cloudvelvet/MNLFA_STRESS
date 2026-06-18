from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUT = Path(r"C:\chen_bauer_2024\한국융합과학회_확인사항_반영문구.docx")


def set_font(run, size=10):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "맑은 고딕")


def add_para(doc, text, size=10, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, size)
    run.bold = bold
    return p


doc = Document()
sec = doc.sections[0]
sec.top_margin = Pt(56)
sec.bottom_margin = Pt(56)
sec.left_margin = Pt(56)
sec.right_margin = Pt(56)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("저자 확인사항 반영 문구")
set_font(run, 13)
run.bold = True

add_para(doc, "방법 절 삽입 문구", 11, True)
add_para(
    doc,
    "연속형 공변량은 분석 전 중심화 또는 표준화하여 투입하였다. 한국어 능력과 가구소득은 표준화 점수로 사용하였고, 연령은 조사 시점에 따른 평균 차이를 줄이기 위해 중심화하여 사용하였다. 성별과 차별 경험은 이분형 변수로 처리하였다. 경험적 선별에서 산출된 p값은 분석에 포함된 문항-공변량 조합 전체에 대해 Benjamini-Hochberg 방식으로 FDR 보정을 적용하였다. 본 연구에서는 FDR < .05이면서 공변량 계수의 절댓값이 .20 이상인 조합을 순서형 DIF 선별 양성 후보로 표시하였다.",
)

add_para(doc, "결과 절 삽입 문구", 11, True)
add_para(
    doc,
    "LLM과 전문가 어휘 기준의 AP 차이에 대해서는 문항 ID를 군집 단위로 한 cluster bootstrap과 paired permutation test를 수행하였다. 1-5차년도 pooled 분석에서 기본 지시문 조건의 AP 차이는 .0066이었으며, 문항 단위 bootstrap 95% 구간은 [-.0700, .0714], permutation test의 양측 p값은 .873이었다. 엄격 지시문 조건의 AP 차이는 .0196이었고, bootstrap 95% 구간은 [-.0625, .0770], permutation test의 양측 p값은 .709였다. 따라서 pooled 분석에서 Gemini 점수가 전문가 어휘 기준을 소폭 상회하였더라도, 이 차이는 통계적으로 안정적인 우위로 해석하기 어렵다.",
)
add_para(
    doc,
    "6차년도 단일 wave 민감도 분석에서는 기본 지시문 조건의 AP 차이가 -.0403이었고, bootstrap 95% 구간은 [-.1223, .0206], permutation test의 양측 p값은 .633이었다. 엄격 지시문 조건의 AP 차이는 -.0560이었으며, bootstrap 95% 구간은 [-.1252, .0208], permutation test의 양측 p값은 .488이었다. 이 결과는 단일 조사시점 기준에서는 오히려 전문가 어휘 기준이 Gemini보다 높은 AP를 보였지만, 그 차이 역시 안정적인 차이로 확인되지는 않았음을 의미한다.",
)

add_para(doc, "표 삽입안", 11, True)
table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"
headers = ["분석", "지시문", "AP 차이", "bootstrap 95% CI", "permutation p"]
for cell, header in zip(table.rows[0].cells, headers):
    p = cell.paragraphs[0]
    r = p.add_run(header)
    set_font(r, 9)
    r.bold = True

rows = [
    ["1-5차 pooled", "기본", ".0066", "[-.0700, .0714]", ".873"],
    ["1-5차 pooled", "엄격", ".0196", "[-.0625, .0770]", ".709"],
    ["6차년도", "기본", "-.0403", "[-.1223, .0206]", ".633"],
    ["6차년도", "엄격", "-.0560", "[-.1252, .0208]", ".488"],
]
for row in rows:
    cells = table.add_row().cells
    for cell, value in zip(cells, row):
        p = cell.paragraphs[0]
        r = p.add_run(value)
        set_font(r, 9)

add_para(
    doc,
    "표 주. AP 차이는 LLM AP에서 전문가 어휘 기준 AP를 뺀 값이다. Bootstrap은 문항 ID를 군집 단위로 하여 500회 반복하였고, permutation test는 동일 문항-공변량 조합 내에서 두 점수의 배정을 교환하는 방식으로 1,000회 반복하였다.",
    9,
)

add_para(doc, "제출 전 메타정보 확인 문구", 11, True)
add_para(
    doc,
    "저자정보, 사사문구, 투고일자는 분석 결과와 무관한 제출용 메타정보이므로 최종 투고 전 별도로 확인하였다.",
)

doc.save(OUT)
print(OUT)
